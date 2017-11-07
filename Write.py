# -*- coding: utf-8 -*-
"""
Created on Fri Nov 03 16:21:55 2017

@author: Administrator
"""

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn


document = Document()
#加入不同等级的标题
chead = 'A股利润表净利润(含少数股东损益)'
ehead = 'A-share profit statement Net profit (including minority gains and losses) (Consolidated)'
ref = 'A股利润表中的净利润(含少数股东损益)字段'
factor = ['LZ_CN_STKA_PRF_COMBO_NET_PRFT_INCL_MIN_INT_INC']
label = ['合并报表']
datatype = 'INT整型'


paragraph = document.add_paragraph()
run = paragraph.add_run(unicode(chead, "utf-8"))
run.font.size = Pt(24)
run.font.name=u'宋体'
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode(ehead, "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'Calibri' 

paragraph = document.add_paragraph()
run = paragraph.add_run(u'一、涉及因子：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')



paragraph = document.add_paragraph()
run = paragraph.add_run(u'1.  数据内容简述：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode(ref, "utf-8"))
run.font.size = Pt(7.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(u'2.  因子列表：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

for i in range(0,len(factor)):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(unicode(label[i]+':', "utf-8"))
    run.font.size = Pt(9)
    paragraph.space_before = Pt(10)
    paragraph.space_after = Pt(10)
    run.font.name=u'宋体' 
    paragraph.paragraph_format.left_indent = Inches(0.75)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')    
    paragraph = document.add_paragraph()
    run = paragraph.add_run(unicode(factor[i], "utf-8"))
    run.font.size = Pt(9)
    run.font.name=u'宋体' 
    paragraph.paragraph_format.left_indent = Inches(0.75)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    paragraph.space_after = Pt(10)

paragraph = document.add_paragraph()
run = paragraph.add_run(u'3.  因子异同：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')



item = '_COMBO_表示：财务报表的合并报表，合并报表以母公司及其子公司组成会计主体\n_PARENT_表示：财务报表的母公司报表，母公司报表以母公司（即上市公司本身）作为会计主体\n不带_Q_表示：常规的季度报、年中报与年报，数据表示从年初至报告期的业绩累加\n带_Q_表示：单季度报表，单季度报表的数据表示季度内的业绩，而非从年初至报告期的业绩累加\n不带_RT表示：静态数据，数据按照每只股票的报表公告日期填写至对应的交易日并向下填充，忽略所有更正报表，只保留更正前的记录，通过根据时间索引整行截断的方式避免未来数据\n带_RT表示：动态数据，数据按照报表的报告期填写，例如年报数据会填写至次年1月1日至3月31日，一季度报数据填写至4月1日至6月30日。根据回测截止时间实时的生成所请求的数据表格，未发布的财报数据会使用上一财报数据进行向下填充。回测截止前公布的更正报表则会覆盖更正前的对应报表数据'

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode(item, "utf-8"))
run.font.size = Pt(7.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(u'二、因子格式：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(u'1.  基本格式：二维类型因子')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')



paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('2.  数据类型：'+datatype, "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('3.  数据单位：无', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('4.  行列索引：', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('1)	行索引：LZ_CN_STKA_EXCH_CAL(A股交易日历)\n', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run = paragraph.add_run(unicode('2)	列索引：LZ_CN_STKA_CODE(A股股票代码)', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')



paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('5.  标识说明：', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('0: 全天停牌或未上市或已退市或未在成分股中\n', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run = paragraph.add_run(unicode('1: 在成分股中且未退市且没有全天停牌', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(u'三、	算法说明：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('无', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.5)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(u'四、	获取与使用代码示例：')
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('1.	获取渠道：PythonSDK、MatlabSDK、读取h5文件\n', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
run = paragraph.add_run(unicode('2.	PythonSDK代码示例\n', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

document.add_picture("python_strategy.jpg", width=Inches(6.5))

paragraph = document.add_paragraph()
run = paragraph.add_run(unicode('3.	MatlabSDK代码示例', "utf-8"))
run.font.size = Pt(10.5)
run.font.name=u'宋体' 
paragraph.paragraph_format.left_indent = Inches(0.75)
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

document.add_picture("matlab_strategy.jpg", width=Inches(6.5))



name = chead+'.docx'
document.save(unicode(name, "utf-8"))
#document.add_heading(u'一级标题',1)
#document.add_heading(u'二级标题',2)
